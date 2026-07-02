from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from google import genai

from starnews.config import Settings

_SECTION_NAMES = ("TITLE", "CAPTION", "HASHTAGS", "SCRIPT")
_SECTION_ALIASES = {
    "TITEL": "TITLE",
    "ÜBERSCHRIFT": "TITLE",
    "CAPTION": "CAPTION",
    "TEASER": "CAPTION",
    "HASHTAGS": "HASHTAGS",
    "HASHTAG": "HASHTAGS",
    "SCRIPT": "SCRIPT",
    "SKRIPT": "SCRIPT",
    "TEXT": "SCRIPT",
}
_HEADER_RE = re.compile(
    r"^\**\s*(?P<name>[A-Za-zÄÖÜäöüß]+)\s*\**\s*:\s*(?P<inline>.*)$"
)
_MIN_SCRIPT_CHARS = 200


@dataclass
class ScriptPackage:
    title: str
    caption: str
    hashtags: str
    script: str

    @property
    def hashtags_list(self) -> list[str]:
        return [tag.strip() for tag in re.findall(r"#\S+", self.hashtags)]


def _normalize_gemini_text(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```[a-zA-Z]*\n?", "", cleaned)
        cleaned = re.sub(r"\n?```$", "", cleaned)
    return cleaned.strip()


def _section_key(raw_name: str) -> str | None:
    name = raw_name.strip().upper()
    name = _SECTION_ALIASES.get(name, name)
    if name in _SECTION_NAMES:
        return name
    return None


def _parse_gemini_output(text: str) -> ScriptPackage:
    sections = {name: "" for name in _SECTION_NAMES}
    current: str | None = None

    for line in _normalize_gemini_text(text).splitlines():
        stripped = line.strip()
        if not stripped:
            if current:
                sections[current] += "\n"
            continue

        match = _HEADER_RE.match(stripped)
        if match:
            key = _section_key(match.group("name"))
            if key:
                current = key
                inline = match.group("inline").strip()
                if inline:
                    sections[current] += inline + "\n"
                continue

        bare = stripped.rstrip(":").strip("*").upper()
        key = _section_key(bare)
        if key and stripped.endswith(":") and not match:
            current = key
            continue

        if current:
            sections[current] += stripped + "\n"

    title = sections["TITLE"].strip()
    caption = sections["CAPTION"].strip()
    hashtags = sections["HASHTAGS"].strip()
    script = sections["SCRIPT"].strip()

    missing = [
        name
        for name, value in (
            ("TITLE", title),
            ("CAPTION", caption),
            ("HASHTAGS", hashtags),
            ("SCRIPT", script),
        )
        if not value
    ]
    if missing:
        preview = _normalize_gemini_text(text)[:400].replace("\n", "\\n")
        raise ValueError(
            "Gemini response missing required sections "
            f"({', '.join(missing)}). Preview: {preview!r}"
        )
    if len(script) < _MIN_SCRIPT_CHARS:
        raise ValueError(
            f"Gemini script too short ({len(script)} chars, need {_MIN_SCRIPT_CHARS}+). "
            "Retry or check the article text."
        )

    return ScriptPackage(title=title, caption=caption, hashtags=hashtags, script=script)


def _response_text(response) -> str:
    text = getattr(response, "text", None)
    if text:
        return text
    parts: list[str] = []
    for candidate in getattr(response, "candidates", None) or []:
        content = getattr(candidate, "content", None)
        for part in getattr(content, "parts", None) or []:
            part_text = getattr(part, "text", None)
            if part_text:
                parts.append(part_text)
    return "\n".join(parts)


def generate_script(article_text: str, settings: Settings) -> ScriptPackage:
    if not settings.gemini_api_key:
        raise ValueError("GEMINI_API_KEY is not set. Add it to ~/.starnews/.env")

    prompt_template = settings.gemini_prompt_file.read_text(encoding="utf-8")
    base_prompt = f"{prompt_template}\n{article_text}"
    retry_suffix = (
        "\n\nWICHTIG: Antworte exakt mit diesen vier Zeilen-Headern, jeweils auf eigener Zeile "
        "und mit Inhalt in den folgenden Zeilen (nicht alles in eine Zeile):\n"
        "TITLE:\nCAPTION:\nHASHTAGS:\nSCRIPT:"
    )

    client = genai.Client(api_key=settings.gemini_api_key)
    last_error: Exception | None = None

    for attempt, prompt in enumerate((base_prompt, base_prompt + retry_suffix)):
        response = client.models.generate_content(
            model=settings.gemini_model,
            contents=prompt,
        )
        text = _response_text(response)
        if not text.strip():
            last_error = ValueError("Gemini returned an empty response")
            continue
        try:
            return _parse_gemini_output(text)
        except ValueError as exc:
            last_error = exc
            if attempt == 0:
                continue
            raise

    raise last_error or ValueError("Gemini script generation failed")


def write_script_docx(package: ScriptPackage, day_dir: Path) -> Path:
    day_dir.mkdir(parents=True, exist_ok=True)
    docx_path = day_dir / "skript.docx"
    doc = Document()
    doc.add_paragraph(package.title)
    doc.add_paragraph(f"Caption:{package.caption} {package.hashtags}")
    doc.add_paragraph(package.script)
    doc.save(docx_path)
    return docx_path
